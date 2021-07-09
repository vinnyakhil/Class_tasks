import os
import docx
import shutil
from docx import Document
def len_bef_num(string):
    length = len(string)
    for i in string:
        if i.isnumeric():
            length = string.find(i)
            break
    return string[:length:].replace('_', '')
def rename_file(name):
    new_name = name.replace('(', '').replace(')', '')
    os.rename(name, new_name)
    return new_name
newdoc = Document()
for i in os.listdir():
    if i.find('.docx') >= 0:
        # rename the file by removing '(' and ')'
        new_name = rename_file(i)
        # get new folder name by taking the file name prior to numeric part and stripping off '-'
        folder = len_bef_num(new_name)
        # Check if folder doesn't exist then create the folder
        if not os.path.isdir(folder):
            os.mkdir(folder)
        # add a heading of level 2 in the new document object
        newdoc.add_heading(new_name, 2)
        # create an instance of a
        # word document we want to open
        doc = Document(new_name)
        # take the text from each para of the source document and add paragraphs and store
        # the object in a variable
        for para in doc.paragraphs:
            doc_para = newdoc.add_paragraph(para.text)
        # add a page break to start a new page
        newdoc.add_page_break()
        shutil.move(os.getcwd() + '\\' + new_name, os.getcwd() + '\\' + folder)
# now save the document to a location
newdoc.save('Master document of all assignments.docx')
